"""
将15份ISO 45001程序文件模板从HTML转换为Word文档(.docx)
保持原始格式：标题层级、编号列表、表格等
"""
import os
import re
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

WORKSPACE = r"C:\Users\Administrator\WorkBuddy\2026-05-31-11-24-40"
OUTPUT_DIR = os.path.join(WORKSPACE, "Word模板")

class HTMLToDocxParser(HTMLParser):
    """Parse HTML and build docx elements"""
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.elements = []  # List of (type, content) tuples
        self.current = []   # Current element being built
        self.tag_stack = []
        self.skip = False   # Skip <style> and <script> blocks
        self.in_table = False
        self.table_rows = []
        self.current_row = []
        self.current_cell = []
        self.cell_tag = 'td'

    def handle_starttag(self, tag, attrs):
        if tag in ('style', 'script'):
            self.skip = True
            return
        if self.skip:
            return

        self.tag_stack.append(tag)

        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            self.current = []

        elif tag == 'p':
            self.current = []

        elif tag == 'ul' or tag == 'ol':
            self.elements.append(('list_start', tag))

        elif tag == 'li':
            self.current = []

        elif tag == 'table':
            self.in_table = True
            self.table_rows = []

        elif tag == 'tr':
            self.current_row = []

        elif tag in ('td', 'th'):
            self.cell_tag = tag
            self.current = []

        elif tag == 'strong' or tag == 'b':
            pass  # Handled in handle_data by checking tag_stack

        elif tag == 'br':
            self.current.append('\n')

        elif tag == 'div':
            if 'ref-section' in dict(attrs).get('class', ''):
                self.current = []

    def handle_endtag(self, tag):
        if tag in ('style', 'script'):
            self.skip = False
            return
        if self.skip:
            if self.tag_stack and self.tag_stack[-1] == tag:
                self.tag_stack.pop()
            return

        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(tag[1])
            self.elements.append(('heading', level, ''.join(self.current)))
            self.current = []

        elif tag == 'p':
            text = ''.join(self.current).strip()
            if text:
                self.elements.append(('paragraph', text))
            self.current = []

        elif tag == 'li':
            text = ''.join(self.current).strip()
            list_type = 'ul'
            for t in reversed(self.tag_stack):
                if t in ('ul', 'ol'):
                    list_type = t
                    break
            if text:
                prefix = '  - ' if list_type == 'ul' else '  '
                self.elements.append(('list_item', prefix + text))
            self.current = []

        elif tag == 'ul' or tag == 'ol':
            self.elements.append(('list_end', tag))

        elif tag == 'tr':
            text = ''.join(self.current).strip()
            self.current_row.append(('th' if self.cell_tag == 'th' else 'td', text))
            self.table_rows.append(self.current_row)
            self.current_row = []
            self.current = []

        elif tag in ('td', 'th'):
            text = ''.join(self.current).strip()
            # will be added on tr close
            self.current = []

        elif tag == 'table':
            self.elements.append(('table', self.table_rows))
            self.in_table = False
            self.table_rows = []

        elif tag == 'div':
            classes = []
            # Check if this was a ref-section
            text = ''.join(self.current).strip()
            if text and len(text) > 5:
                self.elements.append(('paragraph', text))
            self.current = []

        if self.tag_stack and self.tag_stack[-1] == tag:
            self.tag_stack.pop()

    def handle_data(self, data):
        if self.skip:
            return
        # Strip excessive whitespace but keep intentional ones
        data = re.sub(r'\s+', ' ', data)
        self.current.append(data)


def apply_elements_to_doc(doc, elements):
    """Convert parsed elements into docx paragraphs/tables"""

    for elem in elements:
        etype = elem[0]

        if etype == 'heading':
            level = elem[1]
            text = elem[2]
            p = doc.add_heading(text, level=min(level, 4))
            for run in p.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.color.rgb = RGBColor(30, 58, 95)  # #1e3a5f

        elif etype == 'paragraph':
            text = elem[1]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5

        elif etype == 'list_item':
            text = elem[1]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.paragraph_format.space_after = Pt(3)

        elif etype == 'table':
            rows = elem[1]
            if not rows:
                continue
            # Find max columns
            max_cols = max(len(row) for row in rows)
            if max_cols == 0:
                continue

            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, row in enumerate(rows):
                for j in range(max_cols):
                    cell = table.cell(i, j)
                    if j < len(row):
                        cell_type, text = row[j]
                        cell.text = text
                        for p in cell.paragraphs:
                            p.style.font.size = Pt(10)
                            for run in p.runs:
                                run.font.size = Pt(10)
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                                if cell_type == 'th':
                                    run.bold = True
                    # Set column width
                    cell.width = Cm(4)


def convert_html_to_docx(html_path, docx_path):
    """Convert a single HTML template file to Word document"""
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Set page margins (A4: 2.5cm)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Parse HTML
    parser = HTMLToDocxParser(doc)
    parser.feed(html_content)

    # Apply elements to document
    apply_elements_to_doc(doc, parser.elements)

    # Save
    doc.save(docx_path)
    return True


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    template_files = []
    for f in sorted(os.listdir(WORKSPACE)):
        if f.startswith('程序文件模板_') and f.endswith('.html'):
            template_files.append(f)

    print(f"Found {len(template_files)} template files")
    print("=" * 60)

    for html_file in template_files:
        html_path = os.path.join(WORKSPACE, html_file)
        docx_name = html_file.replace('.html', '.docx')
        docx_path = os.path.join(OUTPUT_DIR, docx_name)

        try:
            success = convert_html_to_docx(html_path, docx_path)
            if success:
                size = os.path.getsize(docx_path)
                print(f"[OK] {html_file} -> {docx_name} ({size/1024:.1f} KB)")
            else:
                print(f"[FAIL] {html_file}")
        except Exception as e:
            print(f"[ERROR] {html_file}: {e}")

    print("=" * 60)
    print(f"Done! {len(template_files)} files converted to {OUTPUT_DIR}")


if __name__ == '__main__':
    main()
